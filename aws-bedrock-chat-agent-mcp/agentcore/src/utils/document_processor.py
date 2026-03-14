"""Document processing utilities for SharePoint documents."""

from __future__ import annotations

import logging
from typing import Dict, Any, Optional
from io import BytesIO

logger = logging.getLogger("document_processor")


class DocumentProcessor:
    """Utility class for processing various document types."""
    
    @staticmethod
    def process_document(content_bytes: bytes, file_name: str) -> Dict[str, Any]:
        """
        Process document content and extract text based on file type.
        
        Args:
            content_bytes: Document content as bytes
            file_name: Name of the file (used for type detection)
            
        Returns:
            Dictionary with extracted content and metadata
        """
        file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''
        
        result = {
            "file_name": file_name,
            "file_size": len(content_bytes),
            "file_type": file_extension,
            "content_extracted": False,
        }
        
        try:
            if file_extension in ['txt', 'md', 'json', 'xml', 'csv', 'log']:
                # Plain text files
                result.update(DocumentProcessor._process_text(content_bytes))
            elif file_extension == 'pdf':
                # PDF files
                result.update(DocumentProcessor._process_pdf(content_bytes))
            elif file_extension in ['docx', 'doc']:
                # Word documents
                result.update(DocumentProcessor._process_word(content_bytes))
            elif file_extension in ['xlsx', 'xls']:
                # Excel spreadsheets
                result.update(DocumentProcessor._process_excel(content_bytes))
            else:
                result["message"] = f"Content extraction not supported for .{file_extension} files"
                result["supported_formats"] = ["txt", "md", "json", "xml", "csv", "pdf", "docx", "xlsx"]
        except Exception as e:
            logger.error(f"Error processing document {file_name}: {str(e)}")
            result["extraction_error"] = str(e)
            result["content_extracted"] = False
        
        return result
    
    @staticmethod
    def _process_text(content_bytes: bytes) -> Dict[str, Any]:
        """Process plain text files."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                content = content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                content = content_bytes.decode('latin-1')
            
            return {
                "content": content,
                "content_extracted": True,
                "line_count": len(content.splitlines()),
                "character_count": len(content),
            }
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise
    
    @staticmethod
    def _process_pdf(content_bytes: bytes) -> Dict[str, Any]:
        """Process PDF files."""
        try:
            from PyPDF2 import PdfReader
            
            pdf_reader = PdfReader(BytesIO(content_bytes))
            text_content = []
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            full_text = "\n\n".join(text_content)
            
            return {
                "content": full_text,
                "content_extracted": True,
                "page_count": len(pdf_reader.pages),
                "character_count": len(full_text),
            }
        except ImportError:
            raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
    
    @staticmethod
    def _process_word(content_bytes: bytes) -> Dict[str, Any]:
        """Process Word documents."""
        try:
            from docx import Document
            
            doc = Document(BytesIO(content_bytes))
            
            # Extract paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)
            
            # Combine all text
            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)
            
            return {
                "content": full_text,
                "content_extracted": True,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "character_count": len(full_text),
            }
        except ImportError:
            raise Exception("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            raise
    
    @staticmethod
    def _process_excel(content_bytes: bytes) -> Dict[str, Any]:
        """Process Excel spreadsheets."""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(BytesIO(content_bytes))
            sheets_content = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheets_content.append(f"--- Sheet: {sheet_name} ---")
                sheets_content.append(df.to_string())
                total_rows += len(df)
                total_cols = max(total_cols, len(df.columns))
            
            full_text = "\n\n".join(sheets_content)
            
            return {
                "content": full_text,
                "content_extracted": True,
                "sheet_count": len(excel_file.sheet_names),
                "rows": total_rows,
                "columns": total_cols,
                "character_count": len(full_text),
            }
        except ImportError:
            raise Exception("pandas and openpyxl not installed. Install with: pip install pandas openpyxl")
        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            raise
    
    @staticmethod
    def get_supported_formats() -> Dict[str, str]:
        """
        Get list of supported document formats.
        
        Returns:
            Dictionary mapping file extensions to descriptions
        """
        return {
            "txt": "Plain text files",
            "md": "Markdown files",
            "json": "JSON files",
            "xml": "XML files",
            "csv": "CSV files",
            "log": "Log files",
            "pdf": "PDF documents",
            "docx": "Microsoft Word documents",
            "doc": "Microsoft Word documents (legacy)",
            "xlsx": "Microsoft Excel spreadsheets",
            "xls": "Microsoft Excel spreadsheets (legacy)",
        }
    
    @staticmethod
    def is_supported(file_name: str) -> bool:
        """
        Check if a file type is supported for content extraction.
        
        Args:
            file_name: Name of the file
            
        Returns:
            True if supported, False otherwise
        """
        file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''
        return file_extension in DocumentProcessor.get_supported_formats()